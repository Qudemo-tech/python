#!/usr/bin/env python3
"""
Document Processor Module for Product Knowledge Sources
Extracts text from PDFs, DOCs, and other document formats
"""

import os
import logging
from typing import List, Dict, Optional
import tempfile
from pathlib import Path

# Configure logging
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        """Initialize document processor"""
        self.supported_formats = {
            '.pdf': self._extract_pdf_text,
            '.doc': self._extract_doc_text,
            '.docx': self._extract_docx_text,
            '.txt': self._extract_txt_text,
            '.csv': self._extract_csv_text
        }
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.max_text_length = 100000  # 100k characters
        
    def validate_file(self, file_path: str) -> bool:
        """Validate file size and format"""
        try:
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                logger.warning(f"⚠️ File too large: {file_size} bytes (max: {self.max_file_size})")
                return False
            
            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.supported_formats:
                logger.warning(f"⚠️ Unsupported file format: {file_ext}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"❌ File validation failed: {e}")
            return False
    
    def extract_text(self, file_path: str) -> Optional[Dict]:
        """Extract text from document"""
        try:
            if not self.validate_file(file_path):
                return None
            
            file_ext = Path(file_path).suffix.lower()
            extractor = self.supported_formats.get(file_ext)
            
            if not extractor:
                logger.error(f"❌ No extractor found for format: {file_ext}")
                return None
            
            logger.info(f"📄 Extracting text from: {file_path}")
            text_content = extractor(file_path)
            
            if not text_content or len(text_content.strip()) < 50:
                logger.warning(f"⚠️ No meaningful content extracted from: {file_path}")
                return None
            
            # Truncate if too long
            if len(text_content) > self.max_text_length:
                text_content = text_content[:self.max_text_length] + "..."
                logger.info(f"📝 Text truncated to {self.max_text_length} characters")
            
            return {
                'filename': Path(file_path).name,
                'content': text_content,
                'word_count': len(text_content.split()),
                'source_type': 'document',
                'file_size': os.path.getsize(file_path)
            }
            
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {file_path}: {e}")
            return None
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract page {page_num + 1}: {e}")
                        continue
            
            logger.info(f"✅ Extracted {len(text_content)} characters from PDF")
            return text_content
            
        except ImportError:
            logger.error("❌ PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {e}")
            return ""
    
    def _extract_doc_text(self, file_path: str) -> str:
        """Extract text from DOC file"""
        try:
            import docx2txt
            
            text_content = docx2txt.process(file_path)
            logger.info(f"✅ Extracted {len(text_content)} characters from DOC")
            return text_content
            
        except ImportError:
            logger.error("❌ docx2txt not installed. Install with: pip install docx2txt")
            return ""
        except Exception as e:
            logger.error(f"❌ DOC extraction failed: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_content += row_text + "\n"
            
            logger.info(f"✅ Extracted {len(text_content)} characters from DOCX")
            return text_content
            
        except ImportError:
            logger.error("❌ python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}")
            return ""
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            logger.info(f"✅ Extracted {len(text_content)} characters from TXT")
            return text_content
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text_content = file.read()
                logger.info(f"✅ Extracted {len(text_content)} characters from TXT (latin-1 encoding)")
                return text_content
            except Exception as e:
                logger.error(f"❌ TXT extraction failed with latin-1: {e}")
                return ""
        except Exception as e:
            logger.error(f"❌ TXT extraction failed: {e}")
            return ""
    
    def _extract_csv_text(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            text_content = df.to_string(index=False)
            
            logger.info(f"✅ Extracted {len(text_content)} characters from CSV")
            return text_content
            
        except ImportError:
            logger.error("❌ pandas not installed. Install with: pip install pandas")
            return ""
        except Exception as e:
            logger.error(f"❌ CSV extraction failed: {e}")
            return ""
    
    def chunk_document_content(self, document_data: Dict, chunk_size: int = 1000, overlap: int = 200) -> List[Dict]:
        """Chunk document content for vector storage"""
        try:
            content = document_data['content']
            filename = document_data['filename']
            chunks = []
            
            # Simple text chunking
            pos = 0
            while pos < len(content):
                end = pos + chunk_size
                if end < len(content):
                    # Try to break at sentence boundary
                    for i in range(end, max(pos + chunk_size - 100, pos), -1):
                        if content[i] in '.!?\n':
                            end = i + 1
                            break
                
                chunk_text = content[pos:end].strip()
                if chunk_text:
                    chunks.append({
                        'text': chunk_text,
                        'filename': filename,
                        'source_type': 'document',
                        'start': 0.0,  # No timestamps for documents
                        'end': 0.0
                    })
                
                pos = end - overlap
                if pos >= len(content):
                    break
            
            logger.info(f"📄 Created {len(chunks)} chunks from document: {filename}")
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Document chunking failed: {e}")
            return []
    
    def process_uploaded_file(self, file_content: bytes, filename: str) -> Optional[Dict]:
        """Process uploaded file content"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text from temporary file
                result = self.extract_text(temp_file_path)
                return result
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            logger.error(f"❌ File processing failed: {e}")
            return None
